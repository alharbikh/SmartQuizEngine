import io, re, json, random, urllib.request, urllib.error
from dataclasses import dataclass, asdict
from typing import List, Dict
import streamlit as st
import fitz
from docx import Document

st.set_page_config(page_title='محرك الاختبارات من المستندات', page_icon='🧠', layout='wide')
st.markdown('''<style>html,body,[class*="css"]{direction:rtl;text-align:right}.main .block-container{max-width:1200px;padding-top:1.5rem}h1,h2,h3,p,label,div,span{font-family:Tahoma,Arial,sans-serif}.stButton button{width:100%}textarea,input{direction:rtl!important;text-align:right!important}</style>''', unsafe_allow_html=True)

@dataclass
class Question:
    qid:str; question:str; choices:List[str]; answer:str; explanation:str; source:str; page:int|None=None; qtype:str='اختيار من متعدد'; difficulty:str='متوسط'

def clean_text(t):
    return re.sub(r'\n{3,}','\n\n',re.sub(r'[ \t]+',' ',t.replace('\x00',' '))).strip()

def extract_pdf(b):
    d=fitz.open(stream=b,filetype='pdf'); chunks=[]
    for i,p in enumerate(d,start=1):
        t=clean_text(p.get_text('text'))
        if t: chunks.append(f'[الصفحة {i}]\n{t}')
    return '\n\n'.join(chunks)

def extract_docx(b):
    d=Document(io.BytesIO(b)); blocks=[clean_text(p.text) for p in d.paragraphs if clean_text(p.text)]
    for table in d.tables:
        for row in table.rows:
            vals=[clean_text(c.text) for c in row.cells if clean_text(c.text)]
            if vals: blocks.append(' | '.join(vals))
    return '\n'.join(blocks)

def split_sentences(text):
    raw=re.split(r'(?<=[\.؟!؛:])\s+|\n+',re.sub(r'\[الصفحة \d+\]',' ',text)); return [clean_text(s) for s in raw if 45<=len(clean_text(s))<=500 and len(clean_text(s).split())>=7]

def keywords(s):
    stop={'التي','الذي','الذين','هذه','هذا','ذلك','تلك','هناك','كانت','يكون','يمكن','عندما','حيث','على','إلى','الى','وفي','ومن','كما','فإن','ضمن','بعد','قبل','بين','أكثر','أقل','بشكل','حول','عند','أو','او','هي','هو','تم'}
    return [t for t in re.findall(r'[\u0600-\u06FFA-Za-z0-9_-]{4,}',s) if t not in stop]

def make_mcq(s,pool,idx):
    ks=keywords(s)
    if not ks:return None
    ans=max(ks,key=len); masked=s.replace(ans,'______',1); ds=[]
    for w in pool:
        if w!=ans and w not in ds and 3<=abs(len(w)-len(ans))<=8: ds.append(w)
        if len(ds)==3:break
    ds += ['المفهوم الأساسي','الإجراء المناسب','النتيجة النهائية'][:max(0,3-len(ds))]
    choices=ds[:3]+[ans]; random.shuffle(choices)
    return Question(f'Q{idx:03d}',f'اختر الكلمة أو المصطلح الأنسب لإكمال العبارة التالية:\n{masked}',choices,ans,f'الإجابة مستخرجة من النص: {s}',s)

def generate_rules(text,count):
    ss=split_sentences(text); pool=list(dict.fromkeys(w for s in ss for w in keywords(s))); random.shuffle(ss); out=[]
    for s in ss:
        q=make_mcq(s,pool,len(out)+1)
        if q:out.append(q)
        if len(out)>=count:break
    return [asdict(q) for q in out]

def chunks_with_pages(text,max_chars=6000):
    parts=re.split(r'(?=\[الصفحة \d+\])',text); out=[]
    for p in parts:
        if not p.strip():continue
        m=re.search(r'\[الصفحة (\d+)\]',p); page=int(m.group(1)) if m else None
        body=re.sub(r'^\[الصفحة \d+\]\s*','',p).strip()
        for i in range(0,len(body),max_chars): out.append((page,body[i:i+max_chars]))
    if not out: out=[(None,text[i:i+max_chars]) for i in range(0,len(text),max_chars)]
    return out

def ollama_generate(model,prompt,timeout=180):
    data=json.dumps({'model':model,'prompt':prompt,'stream':False,'format':'json'}).encode()
    req=urllib.request.Request('http://localhost:11434/api/generate',data=data,headers={'Content-Type':'application/json'})
    with urllib.request.urlopen(req,timeout=timeout) as r: return json.loads(r.read().decode())['response']

def ai_questions(text,count,model):
    selected=chunks_with_pages(text)[:max(1,min(8,(count+3)//4))]; per=max(1,(count+len(selected)-1)//len(selected)); allq=[]
    for page,chunk in selected:
        prompt=f'''أنت منشئ اختبارات تعليمية دقيق. استخدم النص المرفق فقط ولا تستخدم معلومات خارجية. أنشئ {per} أسئلة اختيار من متعدد باللغة العربية، لكل سؤال 4 خيارات وإجابة واحدة صحيحة. أعد JSON فقط بالشكل: {{"questions":[{{"question":"...","choices":["...","...","...","..."],"answer":"نص الخيار الصحيح حرفيا","explanation":"تفسير مختصر من النص","source":"اقتباس قصير داعم من النص","difficulty":"سهل|متوسط|صعب"}}]}}. إذا لم يدعم النص سؤالا موثوقا فلا تخترع سؤالا.\n\nالنص:\n{chunk}'''
        obj=json.loads(ollama_generate(model,prompt));
        for item in obj.get('questions',[]):
            if len(item.get('choices',[]))==4 and item.get('answer') in item['choices']:
                item.update({'qid':f'Q{len(allq)+1:03d}','page':page,'qtype':'اختيار من متعدد'}); allq.append(item)
                if len(allq)>=count:return allq
    return allq

def init():
    for k,v in {'extracted_text':'','questions':[]}.items(): st.session_state.setdefault(k,v)
init()
st.title('🧠 محرك الاختبارات من ملفات PDF وWord')
st.caption('استخراج محلي → توليد أسئلة محلي بالذكاء الاصطناعي → مراجعة → اختبار وتصحيح')
with st.sidebar:
    st.header('الإعدادات'); q_count=st.slider('عدد الأسئلة',5,50,10)
    engine=st.radio('محرك إنشاء الأسئلة',['ذكاء اصطناعي محلي (Ollama)','توليد تقليدي بدون AI'])
    model=st.text_input('النموذج المحلي','qwen3:8b',disabled=engine!='ذكاء اصطناعي محلي (Ollama)')
    st.caption('لا يتم إرسال المستند إلى الإنترنت. يحتاج وضع AI إلى Ollama والنموذج مثبتين على الجهاز.')

t1,t2,t3=st.tabs(['1) رفع واستخراج','2) بنك الأسئلة','3) تشغيل الاختبار'])
with t1:
    up=st.file_uploader('اختر ملف PDF أو DOCX',type=['pdf','docx'])
    if up and st.button('استخراج النص',type='primary'):
        try: st.session_state.extracted_text=extract_pdf(up.getvalue()) if up.name.lower().endswith('.pdf') else extract_docx(up.getvalue()); st.session_state.questions=[]
        except Exception as e: st.error(f'تعذر استخراج النص: {e}')
    if st.session_state.extracted_text:
        text=st.session_state.extracted_text; a,b,c=st.columns(3); a.metric('الأحرف',f'{len(text):,}'); b.metric('الكلمات',f'{len(text.split()):,}'); c.metric('المقاطع',len(split_sentences(text))); st.text_area('النص المستخرج',text,height=350)
        if st.button('إنشاء بنك الأسئلة',type='primary'):
            try:
                with st.spinner('جارٍ إنشاء الأسئلة محليًا...'):
                    qs=ai_questions(text,q_count,model) if engine.startswith('ذكاء') else generate_rules(text,q_count)
                st.session_state.questions=qs; st.success(f'تم إنشاء {len(qs)} سؤالًا.') if qs else st.warning('لم يتم إنشاء أسئلة موثوقة من النص.')
            except urllib.error.URLError: st.error('تعذر الاتصال بـ Ollama المحلي. شغّل Ollama ثم نفّذ: ollama pull qwen3:8b')
            except Exception as e: st.error(f'فشل التوليد: {e}')
with t2:
    qs=st.session_state.questions
    if not qs: st.info('أنشئ بنك الأسئلة أولًا.')
    else:
        edited=[]
        for i,q in enumerate(qs):
            with st.expander(f"{q['qid']} — {q['question'][:70]}...",expanded=i==0):
                question=st.text_area('السؤال',q['question'],key=f'q{i}'); choices=[st.text_input(f'الخيار {j+1}',ch,key=f'c{i}{j}') for j,ch in enumerate(q['choices'])]; answer=st.selectbox('الإجابة الصحيحة',choices,index=choices.index(q['answer']) if q['answer'] in choices else 0,key=f'a{i}'); explanation=st.text_area('التفسير',q.get('explanation',''),key=f'e{i}'); st.caption(f"الصفحة: {q.get('page') or 'غير محددة'} | المصدر: {q.get('source','')}"); edited.append({**q,'question':question,'choices':choices,'answer':answer,'explanation':explanation})
        if st.button('حفظ التعديلات'): st.session_state.questions=edited; st.success('تم الحفظ.')
        st.download_button('تنزيل بنك الأسئلة JSON',json.dumps(st.session_state.questions,ensure_ascii=False,indent=2),file_name='question_bank.json',mime='application/json')
with t3:
    qs=st.session_state.questions
    if not qs: st.info('أنشئ بنك الأسئلة أولًا.')
    else:
        answers={}
        for i,q in enumerate(qs,1): st.markdown(f"### {i}. {q['question']}"); answers[q['qid']]=st.radio('اختر الإجابة',q['choices'],index=None,key=f"x{q['qid']}"); st.divider()
        if st.button('تصحيح الاختبار',type='primary'):
            correct=sum(answers.get(q['qid'])==q['answer'] for q in qs); st.metric('النتيجة',f'{correct}/{len(qs)} — {100*correct/len(qs):.0f}%')
            for q in qs:
                if answers.get(q['qid'])!=q['answer']: st.error(f"{q['qid']}: الصحيحة: {q['answer']}"); st.caption(q.get('explanation',''))
st.divider(); st.caption('MVP v0.2 — Local AI via Ollama / Qwen3 8B')
